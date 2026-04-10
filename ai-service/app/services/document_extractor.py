"""
This file is responsible for turning uploaded documents into plain text.

Right now we support PDF and Word files. The rest of the app should not need
to care which parser was used internally; it just asks for extracted text.
"""
from __future__ import annotations

import io
from typing import Optional

from PyPDF2 import PdfReader


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise ValueError(
            "python-docx is required to extract Word documents"
        ) from exc

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def extract_document_text(
    file_bytes: bytes,
    filename: Optional[str] = None,
    content_type: Optional[str] = None,
) -> str:
    """
    Try to extract text from the uploaded file.

    We first look at MIME type because that is usually the most reliable clue.
    If that is missing or unhelpful, we fall back to the filename extension.
    As a last resort, we try both parsers and accept the first one that works.
    """
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if "pdf" in ctype or name.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if (
        "wordprocessingml" in ctype
        or "msword" in ctype
        or name.endswith(".docx")
        or name.endswith(".doc")
    ):
        return _extract_docx(file_bytes)

    # Sometimes uploads come in with a bad content type or no useful file
    # extension. In that case we still try both parsers before giving up.
    try:
        return _extract_pdf(file_bytes)
    except Exception:
        try:
            return _extract_docx(file_bytes)
        except Exception as exc:
            raise ValueError(
                f"Unsupported document format (filename={filename}, content_type={content_type})"
            ) from exc
